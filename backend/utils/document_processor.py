"""
Document Processing Utilities
Handles text extraction from various document formats
"""

import io
import logging
from typing import Optional
from fastapi import UploadFile, HTTPException

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Utility class for processing various document formats"""
    
    @staticmethod
    async def extract_text_from_pdf(content: bytes) -> str:
        """Extract text from PDF document"""
        try:
            import PyPDF2
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
        except ImportError:
            logger.warning("PyPDF2 not installed, attempting basic extraction")
            # Fallback to basic text extraction
            try:
                return content.decode('utf-8', errors='ignore')
            except:
                raise HTTPException(
                    status_code=400,
                    detail="PDF processing requires PyPDF2. Install with: pip install PyPDF2"
                )
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    async def extract_text_from_image(content: bytes) -> str:
        """Extract text from image using OCR"""
        try:
            from PIL import Image
            import pytesseract
            import os
            
            # Configure Tesseract path for Windows
            if os.name == 'nt':  # Windows
                tesseract_paths = [
                    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                    os.path.expanduser(r"~\AppData\Local\Programs\Tesseract-OCR\tesseract.exe")
                ]
                
                for path in tesseract_paths:
                    if os.path.exists(path):
                        pytesseract.pytesseract.tesseract_cmd = path
                        logger.info(f"Using Tesseract OCR from: {path}")
                        break
                else:
                    logger.warning("Tesseract not found in standard locations, using PATH")
            
            image = Image.open(io.BytesIO(content))
            
            # Convert to RGB if necessary
            if image.mode not in ('RGB', 'L'):
                image = image.convert('RGB')
            
            # Perform OCR
            text = pytesseract.image_to_string(image, lang='eng')
            return text.strip()
            
        except ImportError:
            raise HTTPException(
                status_code=400,
                detail="Image OCR requires Pillow and pytesseract. Install with: pip install Pillow pytesseract"
            )
        except pytesseract.TesseractNotFoundError:
            raise HTTPException(
                status_code=400,
                detail="tesseract is not installed or it's not in your PATH. See README file for more information."
            )
        except Exception as e:
            logger.error(f"Error extracting text from image: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Failed to extract text from image: {str(e)}")
    
    @staticmethod
    async def extract_text_from_docx(content: bytes) -> str:
        """Extract text from DOCX document"""
        try:
            import docx
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            
            all_text = "\n".join(paragraphs)
            if tables_text:
                all_text += "\n\nTables:\n" + "\n".join(tables_text)
            
            return all_text.strip()
            
        except ImportError:
            raise HTTPException(
                status_code=400,
                detail="DOCX processing requires python-docx. Install with: pip install python-docx"
            )
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    async def extract_text_from_txt(content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            # Try UTF-8 first
            return content.decode('utf-8').strip()
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                return content.decode('latin-1').strip()
            except:
                # Last resort - ignore errors
                return content.decode('utf-8', errors='ignore').strip()
    
    @classmethod
    async def extract_text(cls, file: UploadFile) -> str:
        """
        Extract text from uploaded file based on file type
        
        Args:
            file: Uploaded file object
            
        Returns:
            Extracted text content
            
        Raises:
            HTTPException: If file type is not supported or extraction fails
        """
        content = await file.read()
        filename_lower = file.filename.lower()
        
        # Route to appropriate extraction method based on file extension
        if filename_lower.endswith('.pdf'):
            return await cls.extract_text_from_pdf(content)
        
        elif filename_lower.endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif')):
            return await cls.extract_text_from_image(content)
        
        elif filename_lower.endswith('.docx'):
            return await cls.extract_text_from_docx(content)
        
        elif filename_lower.endswith(('.txt', '.doc')):
            return await cls.extract_text_from_txt(content)
        
        else:
            # Try to extract as text as a last resort
            try:
                return await cls.extract_text_from_txt(content)
            except:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {file.filename}. Supported: PDF, images, DOCX, TXT"
                )
    
    @staticmethod
    def validate_extracted_text(text: str, min_length: int = 10) -> bool:
        """
        Validate that extracted text is meaningful
        
        Args:
            text: Extracted text
            min_length: Minimum required text length
            
        Returns:
            True if text is valid, False otherwise
        """
        if not text or not text.strip():
            return False
        
        if len(text.strip()) < min_length:
            return False
        
        # Check if text contains mostly printable characters
        printable_ratio = sum(c.isprintable() or c.isspace() for c in text) / len(text)
        return printable_ratio > 0.7
    
    @staticmethod
    def get_supported_formats() -> list:
        """Get list of supported file formats"""
        return [
            ".pdf",
            ".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif",
            ".docx",
            ".txt", ".doc"
        ]
    
    @staticmethod
    def get_format_requirements() -> dict:
        """Get package requirements for each format"""
        return {
            "pdf": "PyPDF2",
            "images": "Pillow, pytesseract (and Tesseract OCR installed on system)",
            "docx": "python-docx",
            "txt": "built-in"
        }

# Create singleton instance
document_processor = DocumentProcessor()
